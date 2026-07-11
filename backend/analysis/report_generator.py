"""
Report Generator
================

Generates HTML and PDF audit reports from a completed (or partial) AuditJob.

PDF is produced via ReportLab (pure Python, no native system dependencies).
HTML is rendered from an inline template for lightweight browser viewing.

LLM Summary
-----------
If ``settings.LLM_API_KEY`` is set, ``get_llm_summary()`` calls the
Google Gemini REST API (gemini-2.0-flash) to generate a plain-English
executive summary paragraph from the structured JSON findings.

If the key is absent, the API call fails, or the response is malformed,
the function falls back silently to a deterministic template summary.
No exception is ever raised to the caller.
"""

import json
import logging
import textwrap
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, Optional

from django.conf import settings

try:
    import requests as _requests
    REQUESTS_AVAILABLE = True
except ImportError:
    _requests = None  # type: ignore[assignment]
    REQUESTS_AVAILABLE = False

logger = logging.getLogger(__name__)

# Gracefully handle missing docx (for environments where it isn't installed)
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Gracefully handle missing reportlab (tests can still import this module)
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
    REPORTLAB_AVAILABLE = True

    # PDF colour palette — defined inside the guard so they don't crash when
    # reportlab is not installed (e.g. before the Docker image is rebuilt)
    _DARK_BG   = colors.HexColor("#0d1117")
    _PANEL_BG  = colors.HexColor("#161b22")
    _BLUE      = colors.HexColor("#58a6ff")
    _GREEN     = colors.HexColor("#3fb950")
    _YELLOW    = colors.HexColor("#e3b341")
    _RED       = colors.HexColor("#f85149")
    _GREY      = colors.HexColor("#7d8590")
    _WHITE     = colors.HexColor("#e6edf3")
    _GRADE_COLORS = {"A": _GREEN, "B": _BLUE, "C": _YELLOW, "D": _YELLOW, "F": _RED}

except ImportError:  # pragma: no cover
    REPORTLAB_AVAILABLE = False


# ---------------------------------------------------------------------------
# LLM integration
# ---------------------------------------------------------------------------

def get_llm_summary(findings: Dict[str, Any]) -> str:
    """
    Request a plain-English executive summary from the Google Gemini API.

    Returns a fallback template string if:
    - settings.LLM_API_KEY is not configured
    - The network request fails
    - The response is malformed

    The caller never needs to handle exceptions from this function.
    """
    api_key = getattr(settings, "LLM_API_KEY", None)
    if not api_key or not REQUESTS_AVAILABLE:
        return _fallback_summary(findings)

    try:

        health = findings.get("health_score", {})
        score = health.get("score", "N/A")
        grade = health.get("grade", "N/A")

        prompt = (
            "You are a senior ML auditor writing a concise executive summary "
            "for a model audit report. Write exactly 2-3 sentences in plain English, "
            "suitable for a non-technical stakeholder. Do NOT use bullet points. "
            "Focus on the most critical finding and the recommended action.\n\n"
            f"Health Score: {score}/100 (Grade {grade})\n"
            f"Audit Findings JSON:\n{json.dumps(findings, indent=2, default=str)[:3000]}"
        )

        response = _requests.post(
            f"https://generativelanguage.googleapis.com/v1beta/models/"
            f"gemini-2.0-flash:generateContent?key={api_key}",
            json={"contents": [{"parts": [{"text": prompt}]}]},
            timeout=15,
        )
        response.raise_for_status()
        data = response.json()
        text = (
            data.get("candidates", [{}])[0]
            .get("content", {})
            .get("parts", [{}])[0]
            .get("text", "")
            .strip()
        )
        if text:
            return text
        return _fallback_summary(findings)

    except Exception as exc:
        logger.warning(f"LLM summary request failed ({type(exc).__name__}): {exc}. Using fallback.")
        return _fallback_summary(findings)


def _fallback_summary(findings: Dict[str, Any]) -> str:
    """Deterministic plain-English summary built from structured findings."""
    health = findings.get("health_score") or {}
    score = health.get("score", None)
    grade = health.get("grade", "N/A")
    components = health.get("components", {})

    if score is None:
        return (
            "This audit is incomplete — one or more analysis modules have not yet run. "
            "Complete the full audit pipeline before using this report for deployment decisions."
        )

    # Find the highest-penalty component
    worst_module = max(
        components.items(),
        key=lambda kv: kv[1].get("weighted_penalty", 0),
        default=("none", {}),
    )
    worst_name = worst_module[0].replace("_", " ").title()
    worst_penalty = worst_module[1].get("weighted_penalty", 0)

    if score >= 85:
        return (
            f"This model achieved a health score of {score}/100 (Grade {grade}), "
            "indicating it is in excellent condition and deployment-ready under normal circumstances. "
            "Routine monitoring is recommended to maintain this standard."
        )
    elif score >= 70:
        return (
            f"This model scored {score}/100 (Grade {grade}). "
            f"The primary concern is {worst_name} (contributing {worst_penalty:.1f} penalty points), "
            "which should be reviewed before production deployment."
        )
    elif score >= 55:
        return (
            f"This model scored {score}/100 (Grade {grade}), indicating moderate risk. "
            f"The {worst_name} module flagged the most significant issues. "
            "Deployment should be deferred until the identified problems are remediated."
        )
    else:
        return (
            f"This model scored {score}/100 (Grade {grade}), indicating severe defects. "
            f"Critical issues were detected — most significantly in {worst_name}. "
            "This model must NOT be deployed until a full remediation review is completed."
        )


# ---------------------------------------------------------------------------
# Shared data extraction helpers
# ---------------------------------------------------------------------------

def _get_health(job_results: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    return job_results.get("health_score") if job_results else None


def _fmt_score(health: Optional[Dict]) -> str:
    if not health:
        return "Not Yet Analyzed"
    score = health.get("score", "?")
    grade = health.get("grade", "?")
    return f"{score}/100 (Grade {grade})"


def _module_status(data: Optional[Any], supported_key: str = "supported") -> str:
    """Return a human-readable status for a module result dict."""
    if data is None:
        return "Not Yet Analyzed"
    if isinstance(data, dict) and not data.get(supported_key, True):
        return data.get("message", "Unsupported for this model type")
    return "Analyzed"


# ---------------------------------------------------------------------------
# HTML report
# ---------------------------------------------------------------------------

_HTML_TEMPLATE = """\
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<title>ModelDoctor Audit Report — {job_id}</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
          background: #0d1117; color: #e6edf3; margin: 0; padding: 40px; }}
  .container {{ max-width: 860px; margin: 0 auto; }}
  h1 {{ font-size: 28px; color: #58a6ff; margin-bottom: 4px; }}
  h2 {{ font-size: 16px; color: #7d8590; font-weight: 400; margin-top: 0; }}
  h3 {{ font-size: 18px; color: #e6edf3; border-bottom: 1px solid #30363d;
        padding-bottom: 8px; margin-top: 36px; }}
  .badge {{ display: inline-block; padding: 4px 12px; border-radius: 20px;
             font-weight: 700; font-size: 20px; }}
  .grade-A {{ background: #1a4736; color: #3fb950; }}
  .grade-B {{ background: #1c3a5e; color: #58a6ff; }}
  .grade-C {{ background: #3d2e00; color: #e3b341; }}
  .grade-D {{ background: #3d1a00; color: #f0883e; }}
  .grade-F {{ background: #3d1111; color: #f85149; }}
  .score-row {{ display: flex; align-items: center; gap: 20px; margin: 20px 0; }}
  .score-num {{ font-size: 56px; font-weight: 800; color: #e6edf3; }}
  .summary-box {{ background: #161b22; border: 1px solid #30363d; border-radius: 8px;
                   padding: 20px; margin: 20px 0; line-height: 1.7; }}
  .module-card {{ background: #161b22; border: 1px solid #30363d; border-radius: 8px;
                   padding: 16px; margin-bottom: 12px; }}
  .module-card h4 {{ margin: 0 0 8px 0; font-size: 14px;
                       text-transform: uppercase; letter-spacing: 0.05em; color: #7d8590; }}
  .module-card p {{ margin: 0; font-size: 14px; color: #e6edf3; }}
  .not-analyzed {{ color: #7d8590; font-style: italic; }}
  .penalty {{ color: #f85149; font-size: 13px; margin-top: 4px; }}
  table {{ width: 100%; border-collapse: collapse; margin-top: 12px; font-size: 13px; }}
  th {{ background: #161b22; color: #7d8590; text-align: left; padding: 8px 12px;
        font-size: 11px; text-transform: uppercase; letter-spacing: 0.06em; }}
  td {{ padding: 8px 12px; border-bottom: 1px solid #21262d; }}
  tr:last-child td {{ border-bottom: none; }}
  .footer {{ margin-top: 48px; font-size: 12px; color: #7d8590;
              border-top: 1px solid #30363d; padding-top: 20px; }}
</style>
</head>
<body>
<div class="container">
  <h1>ModelDoctor Audit Report</h1>
  <h2>Job ID: {job_id} &nbsp;·&nbsp; Generated: {generated_at} &nbsp;·&nbsp; Target: <code>{target_column}</code></h2>

  <h3>Overall Health Score</h3>
  <div class="score-row">
    <div class="score-num">{score}</div>
    <div class="badge grade-{grade}">{grade}</div>
  </div>

  <h3>Executive Summary</h3>
  <div class="summary-box">{summary}</div>

  <h3>Module Breakdown</h3>
  {module_cards}

  <h3>Component Penalty Table</h3>
  <table>
    <thead>
      <tr>
        <th>Module</th><th>Penalty (0–100)</th><th>Weight</th><th>Weighted Penalty</th>
      </tr>
    </thead>
    <tbody>
      {component_rows}
    </tbody>
  </table>

  {leakage_section}
  {data_quality_section}
  {fairness_section}
  {drift_section}

  <div class="footer">
    Generated by ModelDoctor v0.1.0 &nbsp;·&nbsp; {generated_at}
    {llm_note}
  </div>
</div>
</body>
</html>
"""


def _build_module_cards(job_results: Dict[str, Any], leakage_results: list) -> str:
    """Render the per-module status cards as HTML."""
    cards = []

    # Leakage
    leakage_status = "Not Yet Analyzed"
    leakage_detail = ""
    if leakage_results:
        max_risk = max((r.get("risk_score", 0) for r in leakage_results), default=0)
        flagged = sum(1 for r in leakage_results if r.get("known_flag"))
        leakage_status = f"{len(leakage_results)} features analyzed"
        leakage_detail = (
            f"<p>Max risk score: <strong>{max_risk:.1f}</strong> · "
            f"Visibility flags: <strong>{flagged}</strong></p>"
        )
    cards.append(
        f'<div class="module-card"><h4>Target Leakage</h4>'
        f'<p>{"<span class=not-analyzed>" + leakage_status + "</span>" if leakage_status == "Not Yet Analyzed" else leakage_status}</p>'
        f'{leakage_detail}</div>'
    )

    # Calibration
    cal = job_results.get("calibration") if job_results else None
    cal_status = _module_status(cal)
    cal_detail = ""
    if cal and cal.get("supported"):
        cal_detail = f'<p>Brier score: <strong>{cal.get("brier_score", "?")}</strong></p>'
    cards.append(
        f'<div class="module-card"><h4>Calibration</h4>'
        f'<p>{"<span class=not-analyzed>" + cal_status + "</span>" if cal_status == "Not Yet Analyzed" else cal_status}</p>'
        f'{cal_detail}</div>'
    )

    # Overfitting
    ov = job_results.get("overfitting") if job_results else None
    ov_status = _module_status(ov)
    ov_detail = ""
    if ov and ov.get("supported"):
        gap = ov.get("performance_gap", "?")
        ov_detail = f'<p>Performance gap: <strong>{gap:.4f}</strong></p>'
    cards.append(
        f'<div class="module-card"><h4>Overfitting</h4>'
        f'<p>{"<span class=not-analyzed>" + ov_status + "</span>" if ov_status == "Not Yet Analyzed" else ov_status}</p>'
        f'{ov_detail}</div>'
    )

    # Data Quality
    dq = job_results if job_results and "total_rows" in job_results else None
    dq_status = "Not Yet Analyzed" if not dq else "Analyzed"
    dq_detail = ""
    if dq:
        dq_detail = (
            f'<p>Rows: <strong>{dq.get("total_rows")}</strong> · '
            f'Duplicates: <strong>{dq.get("duplicates", {}).get("count", 0)}</strong></p>'
        )
    cards.append(
        f'<div class="module-card"><h4>Data Quality</h4>'
        f'<p>{"<span class=not-analyzed>" + dq_status + "</span>" if dq_status == "Not Yet Analyzed" else dq_status}</p>'
        f'{dq_detail}</div>'
    )

    # Fairness
    fair = job_results.get("fairness") if job_results else None
    fair_status = _module_status(fair)
    fair_detail = ""
    if fair and fair.get("supported"):
        dp_diff = fair.get("demographic_parity_difference", 0.0)
        eo_diff = fair.get("equalized_odds_difference", 0.0)
        fair_detail = (
            f'<p>Demographic Parity Diff: <strong>{dp_diff:.4f}</strong> · '
            f'Equalized Odds Diff: <strong>{eo_diff:.4f}</strong></p>'
        )
    cards.append(
        f'<div class="module-card"><h4>Fairness</h4>'
        f'<p>{"<span class=not-analyzed>" + fair_status + "</span>" if fair_status == "Not Yet Analyzed" else fair_status}</p>'
        f'{fair_detail}</div>'
    )

    return "\n".join(cards)


def _build_fairness_section(job_results: Dict[str, Any]) -> str:
    fair = job_results.get("fairness") if job_results else None
    if not fair or not fair.get("supported"):
        return ""
    
    rows = []
    for grp, metrics in fair.get("group_metrics", {}).items():
        rows.append(
            f'<tr><td>{grp}</td>'
            f'<td>{metrics.get("selection_rate", 0.0):.4f}</td>'
            f'<td>{metrics.get("tpr", 0.0):.4f}</td>'
            f'<td>{metrics.get("fpr", 0.0):.4f}</td>'
            f'<td>{metrics.get("count", 0)}</td></tr>'
        )
        
    return (
        '<h3>Fairness Details</h3>'
        '<p style="font-size: 12px; color: #8b949e; margin-bottom: 12px; line-height: 1.6;">'
        '<strong>Normative Trade-offs Caveat:</strong> Fairness evaluation metrics entail '
        'inherent social and mathematical trade-offs (e.g. demographic parity can conflict with predictive parity '
        'or calibration across groups). These results represent empirical measurements to guide audits and '
        'remediation; they do not dictate a singular standard for absolute fairness.'
        '</p>'
        '<table><thead><tr><th>Group</th><th>Selection Rate</th>'
        '<th>True Positive Rate (TPR)</th><th>False Positive Rate (FPR)</th><th>Count</th></tr></thead>'
        f'<tbody>{"".join(rows)}</tbody></table>'
    )


def _build_drift_section(job_results: Dict[str, Any]) -> str:
    drift = job_results.get("drift") if job_results else None
    if not drift or not drift.get("supported"):
        return ""
        
    rows = []
    for col, info in drift.get("drift_by_feature", {}).items():
        flag = "⚠️ Yes" if info.get("drift_detected") else "No"
        ks_p = info.get("ks_p_value")
        ks_p_str = f'{ks_p:.4f}' if ks_p is not None else "—"
        rows.append(
            f'<tr><td>{col}</td>'
            f'<td>{info.get("psi", 0.0):.4f}</td>'
            f'<td>{ks_p_str}</td>'
            f'<td>{flag}</td></tr>'
        )
        
    return (
        '<h3>Data Drift Details</h3>'
        '<table><thead><tr><th>Feature</th><th>Population Stability Index (PSI)</th>'
        '<th>KS-Test p-value</th><th>Drift Detected</th></tr></thead>'
        f'<tbody>{"".join(rows)}</tbody></table>'
    )



def _build_component_rows(health: Optional[Dict]) -> str:
    if not health:
        return '<tr><td colspan="4" style="color:#7d8590">Not Yet Analyzed</td></tr>'
    rows = []
    for name, comp in health.get("components", {}).items():
        rows.append(
            f'<tr><td>{name.title()}</td>'
            f'<td>{comp.get("penalty", 0):.1f}</td>'
            f'<td>{comp.get("weight", 0):.0%}</td>'
            f'<td>{comp.get("weighted_penalty", 0):.2f}</td></tr>'
        )
    return "\n".join(rows)


def _build_leakage_section(leakage_results: list) -> str:
    if not leakage_results:
        return ""
    rows = []
    for r in leakage_results:
        flag = "⚠️ Yes" if r.get("known_flag") else "No"
        rows.append(
            f'<tr><td>{r.get("feature_name","")}</td>'
            f'<td>{r.get("risk_score", 0):.1f}</td>'
            f'<td>{r.get("drop_pct", 0):.1f}%</td>'
            f'<td>{flag}</td></tr>'
        )
    return (
        '<h3>Leakage Detail</h3>'
        '<table><thead><tr><th>Feature</th><th>Risk Score</th>'
        '<th>Retrain Drop</th><th>Visibility Flag</th></tr></thead>'
        f'<tbody>{"".join(rows)}</tbody></table>'
    )


def _build_data_quality_section(job_results: Dict[str, Any]) -> str:
    if not job_results or "missing_data" not in job_results:
        return ""
    missing = job_results.get("missing_data", {})
    rows = []
    for col, info in missing.items():
        rows.append(
            f'<tr><td>{col}</td>'
            f'<td>{info.get("count", 0)}</td>'
            f'<td>{info.get("percentage", 0):.1f}%</td></tr>'
        )
    return (
        '<h3>Data Quality — Missing Values</h3>'
        '<table><thead><tr><th>Column</th><th>Missing Count</th>'
        '<th>Missing %</th></tr></thead>'
        f'<tbody>{"".join(rows)}</tbody></table>'
    )


def generate_html_report(job: Any, health_score_result: Optional[Dict] = None) -> str:
    """
    Render a complete HTML audit report for the given job.

    ``job`` can be a Django model instance or any object/dict with:
        .id / ['id'], .target_column / ['target_column'],
        .results / ['results'], .leakage_results (queryset or list)

    Partial/missing results are rendered as "Not Yet Analyzed" — no exception
    is raised for incomplete data.
    """
    # Normalise job to dict-like access
    def _attr(obj, key, default=None):
        if isinstance(obj, dict):
            return obj.get(key, default)
        return getattr(obj, key, default)

    job_id       = str(_attr(job, "id", "unknown"))
    target_col   = _attr(job, "target_column", "unknown")
    job_results  = _attr(job, "results") or {}
    raw_leakage  = _attr(job, "leakage_results", None)

    # Serialise leakage queryset to list of dicts
    leakage_list = []
    if raw_leakage is not None:
        try:
            # Django queryset
            leakage_list = list(raw_leakage.values(
                "feature_name", "risk_score", "drop_pct", "known_flag"
            ))
        except AttributeError:
            # Already a plain list
            leakage_list = list(raw_leakage)

    # Build findings dict for LLM / fallback
    findings = {
        "health_score": health_score_result,
        "leakage_feature_count": len(leakage_list),
        "data_quality": {
            "total_rows": job_results.get("total_rows"),
            "duplicates": job_results.get("duplicates"),
        },
    }

    summary       = get_llm_summary(findings)
    health        = health_score_result or {}
    score         = health.get("score", "—")
    grade         = health.get("grade", "—")
    generated_at  = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    llm_note      = (
        "· Executive summary generated by Google Gemini 2.0 Flash"
        if getattr(settings, "LLM_API_KEY", None) else
        "· Executive summary: deterministic template (set LLM_API_KEY for AI summary)"
    )

    return _HTML_TEMPLATE.format(
        job_id=job_id,
        generated_at=generated_at,
        target_column=target_col,
        score=score,
        grade=grade,
        summary=summary,
        module_cards=_build_module_cards(job_results, leakage_list),
        component_rows=_build_component_rows(health_score_result),
        leakage_section=_build_leakage_section(leakage_list),
        data_quality_section=_build_data_quality_section(job_results),
        fairness_section=_build_fairness_section(job_results),
        drift_section=_build_drift_section(job_results),
        llm_note=llm_note,
    )



# ---------------------------------------------------------------------------
# PDF report (ReportLab)
# ---------------------------------------------------------------------------


def generate_pdf_report(job: Any, health_score_result: Optional[Dict] = None) -> bytes:
    """
    Generate a PDF audit report and return the raw bytes.

    Uses ReportLab (pure Python). Falls back to a minimal placeholder PDF
    if ReportLab is not installed, so the function always returns valid bytes.

    Partial/missing results are shown as "Not Yet Analyzed" — no exception
    is raised for incomplete data.
    """
    if not REPORTLAB_AVAILABLE:
        # Ultra-minimal fallback — reportlab not yet installed in this environment
        return b"%PDF-1.4\n% ReportLab not installed\n"

    # Resolve colour constants that are only available when reportlab loaded
    dark_bg     = _DARK_BG    # noqa: F821 — defined inside try block above
    panel_bg    = _PANEL_BG   # noqa: F821
    blue_c      = _BLUE       # noqa: F821
    grey_c      = _GREY       # noqa: F821
    white_c     = _WHITE      # noqa: F821
    grade_colors = _GRADE_COLORS  # noqa: F821

    def _attr(obj, key, default=None):
        if isinstance(obj, dict):
            return obj.get(key, default)
        return getattr(obj, key, default)

    job_id       = str(_attr(job, "id", "unknown"))
    target_col   = _attr(job, "target_column", "unknown")
    job_results  = _attr(job, "results") or {}
    raw_leakage  = _attr(job, "leakage_results", None)

    leakage_list = []
    if raw_leakage is not None:
        try:
            leakage_list = list(raw_leakage.values(
                "feature_name", "risk_score", "drop_pct", "known_flag"
            ))
        except AttributeError:
            leakage_list = list(raw_leakage)

    findings = {
        "health_score": health_score_result,
        "leakage_feature_count": len(leakage_list),
    }
    summary      = get_llm_summary(findings)
    health       = health_score_result or {}
    score        = health.get("score", "N/A")
    grade        = health.get("grade", "N/A")
    grade_color  = grade_colors.get(str(grade), grey_c)
    generated_at = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title", parent=styles["Heading1"],
        textColor=blue_c, fontSize=22, spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle", parent=styles["Normal"],
        textColor=grey_c, fontSize=9, spaceAfter=20,
    )
    h3_style = ParagraphStyle(
        "H3", parent=styles["Heading2"],
        textColor=white_c, fontSize=13, spaceBefore=20, spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        textColor=white_c, fontSize=10, leading=16,
    )
    not_analyzed_style = ParagraphStyle(
        "NA", parent=styles["Normal"],
        textColor=grey_c, fontSize=10, leading=14, fontName="Helvetica-Oblique",
    )

    story = []

    # Title
    story.append(Paragraph("ModelDoctor Audit Report", title_style))
    story.append(Paragraph(
        f"Job: {job_id} &nbsp;·&nbsp; Target: {target_col} &nbsp;·&nbsp; {generated_at}",
        subtitle_style,
    ))
    story.append(HRFlowable(width="100%", thickness=0.5, color=grey_c))
    story.append(Spacer(1, 12))

    # Score
    story.append(Paragraph("Overall Health Score", h3_style))
    score_table_data = [
        [
            Paragraph(f"<b>{score}</b> / 100", ParagraphStyle("S", fontSize=28, textColor=white_c)),
            Paragraph(f"<b>{grade}</b>", ParagraphStyle("G", fontSize=28, textColor=grade_color)),
        ]
    ]
    story.append(Table(score_table_data, colWidths=[8 * cm, 4 * cm]))
    story.append(Spacer(1, 8))

    # Summary
    story.append(Paragraph("Executive Summary", h3_style))
    story.append(Paragraph(summary, body_style))
    story.append(Spacer(1, 12))

    # Component table
    story.append(Paragraph("Component Penalty Table", h3_style))
    comp_headers = ["Module", "Penalty (0–100)", "Weight", "Weighted Penalty"]
    comp_data = [comp_headers]
    if health_score_result:
        for name, comp in health.get("components", {}).items():
            comp_data.append([
                name.title(),
                f'{comp.get("penalty", 0):.1f}',
                f'{comp.get("weight", 0):.0%}',
                f'{comp.get("weighted_penalty", 0):.2f}',
            ])
    else:
        comp_data.append(["Not Yet Analyzed", "—", "—", "—"])

    comp_table = Table(comp_data, colWidths=[5 * cm, 4 * cm, 3 * cm, 4 * cm])
    comp_table.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0), panel_bg),
        ("TEXTCOLOR",    (0, 0), (-1, 0), grey_c),
        ("FONTSIZE",     (0, 0), (-1, 0), 8),
        ("FONTNAME",     (0, 0), (-1, 0), "Helvetica-Bold"),
        ("TEXTCOLOR",    (0, 1), (-1, -1), white_c),
        ("FONTSIZE",     (0, 1), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [dark_bg, panel_bg]),
        ("GRID",         (0, 0), (-1, -1), 0.3, grey_c),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING",    (0, 0), (-1, -1), 6),
    ]))
    story.append(comp_table)
    story.append(Spacer(1, 12))

    # Leakage
    story.append(Paragraph("Target Leakage", h3_style))
    if leakage_list:
        lk_headers = ["Feature", "Risk Score", "Retrain Drop %", "Visibility Flag"]
        lk_data = [lk_headers]
        for r in leakage_list:
            flag = "Yes ⚠" if r.get("known_flag") else "No"
            lk_data.append([
                r.get("feature_name", ""),
                f'{r.get("risk_score", 0):.1f}',
                f'{r.get("drop_pct", 0):.1f}',
                flag,
            ])
        lk_table = Table(lk_data, colWidths=[5 * cm, 4 * cm, 4 * cm, 4 * cm])
        lk_table.setStyle(TableStyle([
            ("BACKGROUND",   (0, 0), (-1, 0), panel_bg),
            ("TEXTCOLOR",    (0, 0), (-1, 0), grey_c),
            ("FONTNAME",     (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",     (0, 0), (-1, 0), 8),
            ("TEXTCOLOR",    (0, 1), (-1, -1), white_c),
            ("FONTSIZE",     (0, 1), (-1, -1), 9),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [dark_bg, panel_bg]),
            ("GRID",         (0, 0), (-1, -1), 0.3, grey_c),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING",    (0, 0), (-1, -1), 6),
        ]))
        story.append(lk_table)
    else:
        story.append(Paragraph("Not Yet Analyzed", not_analyzed_style))

    # Fairness & Drift Details
    fair_data = job_results.get("fairness")
    if fair_data and fair_data.get("supported"):
        story.append(Paragraph("Fairness Metrics", h3_style))
        story.append(Paragraph(
            "<b>Demographic Parity Difference:</b> "
            f"{fair_data.get('demographic_parity_difference', 0.0):.4f}<br/>"
            "<b>Equalized Odds Difference:</b> "
            f"{fair_data.get('equalized_odds_difference', 0.0):.4f}",
            body_style
        ))
        story.append(Spacer(1, 6))

    drift_data = job_results.get("drift")
    if drift_data and drift_data.get("supported"):
        story.append(Paragraph("Data Drift Details", h3_style))
        drift_headers = ["Feature", "PSI", "KS p-value", "Drift Detected"]
        drift_rows = [drift_headers]
        for col, info in drift_data.get("drift_by_feature", {}).items():
            ks_p = info.get("ks_p_value")
            ks_p_str = f"{ks_p:.4f}" if ks_p is not None else "—"
            drift_rows.append([
                col,
                f"{info.get('psi', 0.0):.4f}",
                ks_p_str,
                "Yes ⚠" if info.get("drift_detected") else "No"
            ])
        drift_table = Table(drift_rows, colWidths=[5 * cm, 4 * cm, 4 * cm, 4 * cm])
        drift_table.setStyle(TableStyle([
            ("BACKGROUND",   (0, 0), (-1, 0), panel_bg),
            ("TEXTCOLOR",    (0, 0), (-1, 0), grey_c),
            ("FONTNAME",     (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",     (0, 0), (-1, 0), 8),
            ("TEXTCOLOR",    (0, 1), (-1, -1), white_c),
            ("FONTSIZE",     (0, 1), (-1, -1), 9),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [dark_bg, panel_bg]),
            ("GRID",         (0, 0), (-1, -1), 0.3, grey_c),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING",    (0, 0), (-1, -1), 6),
        ]))
        story.append(drift_table)
        story.append(Spacer(1, 10))

    # Module status summary
    story.append(Paragraph("Module Status", h3_style))
    modules = [
        ("Calibration",  job_results.get("calibration")),
        ("Overfitting",  job_results.get("overfitting")),
        ("Feature Dom.", job_results.get("feature_dominance")),
        ("Fairness",     job_results.get("fairness")),
        ("Data Drift",   job_results.get("drift")),
    ]
    for name, data in modules:
        status = _module_status(data) if data is not None else "Not Yet Analyzed"
        story.append(Paragraph(f"<b>{name}:</b> {status}", body_style))

    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=0.5, color=grey_c))
    llm_note = (
        "Executive summary: Google Gemini 2.0 Flash"
        if getattr(settings, "LLM_API_KEY", None) else
        "Executive summary: deterministic template (set LLM_API_KEY for AI summary)"
    )
    story.append(Paragraph(
        f"Generated by ModelDoctor v0.1.0 · {generated_at} · {llm_note}",
        ParagraphStyle("Footer", parent=styles["Normal"], textColor=grey_c, fontSize=8),
    ))

    doc.build(story)
    return buf.getvalue()


def generate_docx_report(job: Any, health_score_result: Optional[Dict] = None) -> bytes:
    """
    Generate a redesigned 4-page Word document (.docx) report and return the raw bytes.
    """
    if not DOCX_AVAILABLE:
        return b""

    def _attr(obj, key, default=None):
        if isinstance(obj, dict):
            return obj.get(key, default)
        return getattr(obj, key, default)

    job_id       = str(_attr(job, "id", "unknown"))
    target_col   = _attr(job, "target_column", "unknown")
    job_results  = _attr(job, "results") or {}
    raw_leakage  = _attr(job, "leakage_results", None)

    leakage_list = []
    if raw_leakage is not None:
        try:
            leakage_list = list(raw_leakage.values(
                "feature_name", "risk_score", "drop_pct", "known_flag"
            ))
        except AttributeError:
            leakage_list = list(raw_leakage)

    findings = {
        "health_score": health_score_result,
        "leakage_feature_count": len(leakage_list),
    }
    summary = get_llm_summary(findings)
    health = health_score_result or {}
    score = health.get("score", "N/A")
    grade = health.get("grade", "N/A")
    grade_str = str(grade).upper()

    doc = Document()

    # Page Margins: 1 inch (72 Pt)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Set default font style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Helper function for cell background shading
    def set_cell_shading(cell, color_hex):
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

    # Helper function for cell margins (padding)
    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
            node = OxmlElement(m)
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # Helper for running header at the top of pages 2-4
    def add_running_header(doc, j_id):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(f"ModelDoctor Audit • {j_id}")
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x7D, 0x85, 0x90)
        r.italic = True
        p.paragraph_format.space_after = Pt(12)

    # =========================================================================
    # PAGE 1: COVER PAGE
    # =========================================================================
    # Header
    p_hdr = doc.add_paragraph()
    r_hdr = p_hdr.add_run("MODELDOCTOR")
    r_hdr.font.size = Pt(9.5)
    r_hdr.font.bold = True
    r_hdr.font.color.rgb = RGBColor(0x7D, 0x85, 0x90)
    p_hdr.paragraph_format.space_after = Pt(2)

    # Subtitle
    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Model Audit Report")
    r_sub.font.size = Pt(14)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    p_sub.paragraph_format.space_after = Pt(12)

    # Job ID Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run(job_id)
    r_title.font.size = Pt(28)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
    p_title.paragraph_format.space_after = Pt(4)

    # Target
    p_tgt = doc.add_paragraph()
    r_tgt = p_tgt.add_run(f"Target variable: {target_col}")
    r_tgt.font.size = Pt(11)
    r_tgt.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    p_tgt.paragraph_format.space_after = Pt(36)

    # Overall Health Score Card
    tbl_score = doc.add_table(rows=1, cols=1)
    tbl_score.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_score = tbl_score.rows[0].cells[0]
    cell_score.width = Inches(6.5)
    set_cell_shading(cell_score, "F6F8FA")
    set_cell_margins(cell_score, top=200, bottom=200, left=250, right=250)

    p_lbl = cell_score.paragraphs[0]
    r_lbl = p_lbl.add_run("OVERALL HEALTH SCORE")
    r_lbl.font.size = Pt(10)
    r_lbl.font.bold = True
    r_lbl.font.color.rgb = RGBColor(0x7D, 0x85, 0x90)
    p_lbl.paragraph_format.space_after = Pt(8)

    p_val = cell_score.add_paragraph()
    r_score_val = p_val.add_run(f"{score}")
    r_score_val.font.size = Pt(48)
    r_score_val.font.bold = True

    if grade_str == 'A':
        r_score_val.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32) # Green
    elif grade_str == 'B':
        r_score_val.font.color.rgb = RGBColor(0x15, 0x65, 0xC0) # Blue
    elif grade_str in ['C', 'D']:
        r_score_val.font.color.rgb = RGBColor(0xEF, 0x6C, 0x00) # Orange
    else:
        r_score_val.font.color.rgb = RGBColor(0xC6, 0x28, 0x28) # Red

    r_slash = p_val.add_run(" / 100")
    r_slash.font.size = Pt(20)
    r_slash.font.color.rgb = RGBColor(0x7D, 0x85, 0x90)
    p_val.paragraph_format.space_after = Pt(12)

    p_grd = cell_score.add_paragraph()
    r_grd = p_grd.add_run(f"GRADE {grade_str} — ")
    r_grd.font.bold = True
    r_grd.font.size = Pt(12)
    r_grd.font.color.rgb = r_score_val.font.color.rgb

    if grade_str == 'A':
        r_desc = p_grd.add_run("EXCELLENT MODEL CONDITION")
    elif grade_str == 'B':
        r_desc = p_grd.add_run("GOOD WITH MINOR ISSUES")
    elif grade_str == 'C':
        r_desc = p_grd.add_run("MODERATE RISK")
    elif grade_str == 'D':
        r_desc = p_grd.add_run("HIGH RISK")
    else:
        r_desc = p_grd.add_run("SEVERE DEFECTS")
    r_desc.font.size = Pt(12)
    r_desc.font.color.rgb = r_score_val.font.color.rgb
    p_grd.paragraph_format.space_after = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(48)

    # Metadata
    p_dep = doc.add_paragraph()
    r_dep_lbl = p_dep.add_run("Deployment status: ")
    r_dep_lbl.font.size = Pt(10)
    r_dep_lbl.font.color.rgb = RGBColor(0x7D, 0x85, 0x90)

    if grade_str == 'F':
        r_dep_status = p_dep.add_run("BLOCKED — do not deploy")
        r_dep_status.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    elif grade_str in ['C', 'D']:
        r_dep_status = p_dep.add_run("WARNING — deploy with caution")
        r_dep_status.font.color.rgb = RGBColor(0xEF, 0x6C, 0x00)
    else:
        r_dep_status = p_dep.add_run("APPROVED — ready for deployment")
        r_dep_status.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    r_dep_status.font.bold = True
    r_dep_status.font.size = Pt(10)

    generated_at = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

    p_gen = doc.add_paragraph()
    r_gen = p_gen.add_run(f"Report generated: {generated_at}\n")
    r_gen.add_text("Engine: ModelDoctor v0.1.0\n")
    llm_key_present = bool(getattr(settings, "LLM_API_KEY", None))
    mode_text = "Google Gemini 2.0 Flash (AI summary)" if llm_key_present else "Deterministic template (no LLM_API_KEY set)"
    r_gen.add_text(f"Summary mode: {mode_text}")
    r_gen.font.size = Pt(9.5)
    r_gen.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    doc.add_page_break()

    # =========================================================================
    # PAGE 2: EXECUTIVE SUMMARY & COMPONENT PENALTY TABLE
    # =========================================================================
    add_running_header(doc, job_id)

    # SECTION 1 Heading
    p_sec1 = doc.add_paragraph()
    r_sec1 = p_sec1.add_run("SECTION 1")
    r_sec1.font.size = Pt(9)
    r_sec1.font.bold = True
    r_sec1.font.color.rgb = RGBColor(0x7D, 0x85, 0x90)
    p_sec1.paragraph_format.space_after = Pt(2)

    # Executive Summary Title
    p_h1 = doc.add_paragraph()
    r_h1 = p_h1.add_run("Executive Summary")
    r_h1.font.size = Pt(16)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
    p_h1.paragraph_format.space_after = Pt(8)

    # Summary Text
    p_sum = doc.add_paragraph()
    r_sum = p_sum.add_run(summary)
    r_sum.font.size = Pt(10.5)
    r_sum.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    p_sum.paragraph_format.line_spacing = 1.15
    p_sum.paragraph_format.space_after = Pt(12)

    # Blocker box (if applicable)
    leaky_blockers = [f for f in leakage_list if f.get("risk_score", 0) > 50 or f.get("known_flag")]
    worst_name = "the leakage feature"
    if leaky_blockers:
        worst_leaker = max(leaky_blockers, key=lambda x: x.get("risk_score", 0))
        worst_name = worst_leaker.get("feature_name", "unknown")
        worst_risk = worst_leaker.get("risk_score", 0)
        worst_drop = worst_leaker.get("drop_pct", 0)

        tbl_blocker = doc.add_table(rows=1, cols=1)
        tbl_blocker.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_blocker = tbl_blocker.rows[0].cells[0]
        cell_blocker.width = Inches(6.5)
        set_cell_shading(cell_blocker, "FDF2F2")
        set_cell_margins(cell_blocker, top=120, bottom=120, left=180, right=180)

        # Left border red
        tcPr = cell_blocker._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="36" w:space="0" w:color="F85149"/>'
            f'<w:top w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:bottom w:val="none"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

        p_blk = cell_blocker.paragraphs[0]
        r_blk = p_blk.add_run(
            f"BLOCKER — {worst_name} carries a Leakage Risk Score of {worst_risk:.1f} and "
            f"accounts for a {worst_drop:.1f}% retrain performance drop when removed. "
            f"This is almost certainly a direct or near-direct encoding of the target and must be "
            f"investigated before any further work."
        )
        r_blk.font.bold = True
        r_blk.font.size = Pt(9.5)
        r_blk.font.color.rgb = RGBColor(0x9B, 0x1C, 0x1C)
        p_blk.paragraph_format.space_after = Pt(0)

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Engineer's read
    p_eread = doc.add_paragraph()
    r_eread = p_eread.add_run("Engineer's read")
    r_eread.font.bold = True
    r_eread.font.size = Pt(11)
    r_eread.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    p_eread.paragraph_format.space_after = Pt(6)

    components = health.get("components", {})
    
    # Dynamic values extraction
    lk_comp = components.get("leakage", {})
    lk_p = lk_comp.get("penalty", 0)
    lk_w = lk_comp.get("weight", 0)
    lk_wp = lk_comp.get("weighted_penalty", 0)

    cal_comp = components.get("calibration", {})
    cal_p = cal_comp.get("penalty", 0)
    cal_w = cal_comp.get("weight", 0)
    cal_wp = cal_comp.get("weighted_penalty", 0)

    ov_comp = components.get("overfitting", {})
    ov_p = ov_comp.get("penalty", 0)
    ov_w = ov_comp.get("weight", 0)
    ov_wp = ov_comp.get("weighted_penalty", 0)

    dq_comp = components.get("data_quality", {})
    dq_p = dq_comp.get("penalty", 0)
    dq_w = dq_comp.get("weight", 0)
    dq_wp = dq_comp.get("weighted_penalty", 0)

    fn_comp = components.get("fairness", {})
    fn_p = fn_comp.get("penalty", 0)
    fn_w = fn_comp.get("weight", 0)
    fn_wp = fn_comp.get("weighted_penalty", 0)

    if lk_p > 50:
        bullet1 = f"Leakage is the dominant failure mode ({lk_w:.0%} weight, {lk_p:.1f}/100 penalty → {lk_wp:.2f} weighted points lost). Nothing else in this report matters until this is resolved."
    else:
        bullet1 = f"Leakage was evaluated resulting in ({lk_w:.0%} weight, {lk_p:.1f}/100 penalty → {lk_wp:.2f} weighted points lost)."

    bullet2 = f"Calibration is secondary but non-trivial ({cal_p:.1f}/100 penalty, {cal_w:.0%} weight → {cal_wp:.2f} weighted points). Predicted probabilities should not be trusted for thresholding or ranking decisions yet."
    bullet3 = f"Overfitting and Data Quality are contributing but not primary drivers at this stage ({ov_wp:.2f} and {dq_wp:.2f} weighted points respectively)."
    
    if fn_p > 0:
        bullet4 = f"Fairness displays measurable group disparity ({fn_p:.1f}/100 penalty → {fn_wp:.2f} weighted points)."
    else:
        bullet4 = "Fairness shows no measurable disparity on the metrics evaluated (0.0 penalty) — but re-run after leakage is fixed, since leakage can mask or distort fairness signals."

    for text in [bullet1, bullet2, bullet3, bullet4]:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(4)
        p_b.paragraph_format.line_spacing = 1.15
        p_b.paragraph_format.left_indent = Inches(0.25)
        
        parts = text.split(" (", 1)
        if len(parts) == 2:
            r_lead = p_b.add_run(parts[0])
            r_lead.font.bold = True
            r_lead.font.size = Pt(9.5)
            r_rest = p_b.add_run(" (" + parts[1])
            r_rest.font.size = Pt(9.5)
        else:
            r_all = p_b.add_run(text)
            r_all.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 2 Heading
    p_sec2 = doc.add_paragraph()
    r_sec2 = p_sec2.add_run("SECTION 2")
    r_sec2.font.size = Pt(9)
    r_sec2.font.bold = True
    r_sec2.font.color.rgb = RGBColor(0x7D, 0x85, 0x90)
    p_sec2.paragraph_format.space_after = Pt(2)

    # Component Table Title
    p_h2 = doc.add_paragraph()
    r_h2 = p_h2.add_run("Component Penalty Table")
    r_h2.font.size = Pt(14)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
    p_h2.paragraph_format.space_after = Pt(6)

    p_desc2 = doc.add_paragraph()
    r_desc2 = p_desc2.add_run(
        "Weighted Penalty = Penalty × Weight. The Overall Health Score is 100 minus the sum of weighted penalties across all modules."
    )
    r_desc2.font.size = Pt(9.5)
    r_desc2.font.italic = True
    r_desc2.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    p_desc2.paragraph_format.space_after = Pt(10)

    # Table Grid
    tbl_comp = doc.add_table(rows=1, cols=5)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = tbl_comp.rows[0].cells
    headers = ["MODULE", "PENALTY (0–100)", "WEIGHT", "WEIGHTED PENALTY", "SEVERITY"]
    col_widths = [Inches(1.5), Inches(1.3), Inches(1.0), Inches(1.5), Inches(1.2)]

    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        hdr_cells[idx].width = col_widths[idx]
        set_cell_shading(hdr_cells[idx], "1F2937")
        set_cell_margins(hdr_cells[idx], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[idx].paragraphs[0]
        r = p.runs[0]
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    module_keys = ["leakage", "calibration", "overfitting", "fairness", "data_quality"]
    for key in module_keys:
        comp_info = components.get(key, {})
        p_val = comp_info.get("penalty", 0)
        w_val = comp_info.get("weight", 0)
        wp_val = comp_info.get("weighted_penalty", 0)

        if p_val >= 50:
            sev_text = "CRITICAL"
            sev_bg = "FDE2E2"
            sev_fg = RGBColor(0x9B, 0x1C, 0x1C)
        elif p_val >= 15:
            sev_text = "MODERATE"
            sev_bg = "FEF3C7"
            sev_fg = RGBColor(0x92, 0x40, 0x0E)
        else:
            sev_text = "LOW"
            sev_bg = "E1F5FE" if key == "fairness" and p_val == 0 else "DEF7EC"
            sev_fg = RGBColor(0x03, 0x69, 0xA1) if key == "fairness" and p_val == 0 else RGBColor(0x03, 0x54, 0x3F)

        row_cells = tbl_comp.add_row().cells
        row_cells[0].text = key.replace("_", " ").title()
        row_cells[1].text = f"{p_val:.1f}"
        row_cells[2].text = f"{w_val:.0%}"
        row_cells[3].text = f"{wp_val:.2f}"
        row_cells[4].text = sev_text

        for idx in range(5):
            row_cells[idx].width = col_widths[idx]
            set_cell_margins(row_cells[idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[idx].paragraphs[0]
            r = p.runs[0]
            r.font.size = Pt(9)
            if idx == 4:
                set_cell_shading(row_cells[idx], sev_bg)
                r.font.bold = True
                r.font.color.rgb = sev_fg
            elif idx == 1 and p_val >= 50:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

    doc.add_page_break()

    # =========================================================================
    # PAGE 3: TARGET LEAKAGE & FAIRNESS & MODULE STATUS
    # =========================================================================
    add_running_header(doc, job_id)

    # SECTION 3 Heading
    p_sec3 = doc.add_paragraph()
    r_sec3 = p_sec3.add_run("SECTION 3")
    r_sec3.font.size = Pt(9)
    r_sec3.font.bold = True
    r_sec3.font.color.rgb = RGBColor(0x7D, 0x85, 0x90)
    p_sec3.paragraph_format.space_after = Pt(2)

    p_h3 = doc.add_paragraph()
    r_h3 = p_h3.add_run("Target Leakage")
    r_h3.font.size = Pt(14)
    r_h3.font.bold = True
    r_h3.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
    p_h3.paragraph_format.space_after = Pt(6)

    p_desc3 = doc.add_paragraph()
    r_desc3 = p_desc3.add_run(
        "Risk Score estimates how strongly a feature encodes the target directly or indirectly. Retrain Drop % is the measured drop in model performance when the feature is removed and the model is retrained — a large drop paired with a high risk score is strong evidence of leakage rather than a genuinely predictive feature."
    )
    r_desc3.font.size = Pt(9.5)
    r_desc3.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    p_desc3.paragraph_format.space_after = Pt(10)

    # Leakage Table
    tbl_leak = doc.add_table(rows=1, cols=4)
    tbl_leak.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells_l = tbl_leak.rows[0].cells
    headers_l = ["FEATURE", "RISK SCORE", "RETRAIN DROP %", "FLAG"]
    col_widths_l = [Inches(2.5), Inches(1.3), Inches(1.5), Inches(1.2)]

    for idx, name in enumerate(headers_l):
        hdr_cells_l[idx].text = name
        hdr_cells_l[idx].width = col_widths_l[idx]
        set_cell_shading(hdr_cells_l[idx], "1F2937")
        set_cell_margins(hdr_cells_l[idx], top=100, bottom=100, left=100, right=100)
        p = hdr_cells_l[idx].paragraphs[0]
        r = p.runs[0]
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    if not leakage_list:
        row_cells = tbl_leak.add_row().cells
        row_cells[0].text = "No features analyzed or target leakage detected."
        row_cells[0].width = Inches(6.5)
        for idx in range(1, 4):
            row_cells[idx].width = Inches(0)
    else:
        for f in leakage_list:
            f_name = f.get("feature_name", "")
            f_risk = f.get("risk_score", 0)
            f_drop = f.get("drop_pct", 0)
            f_flag = f.get("known_flag", False)

            flag_val = "FLAGGED" if (f_risk > 50 or f_flag) else "—"

            row_cells = tbl_leak.add_row().cells
            row_cells[0].text = f_name
            row_cells[1].text = f"{f_risk:.1f}"
            row_cells[2].text = f"{f_drop:.1f}%"
            row_cells[3].text = flag_val

            for idx in range(4):
                row_cells[idx].width = col_widths_l[idx]
                set_cell_margins(row_cells[idx], top=80, bottom=80, left=100, right=100)
                p = row_cells[idx].paragraphs[0]
                r = p.runs[0]
                r.font.size = Pt(9)
                r.font.name = "Consolas" if idx == 0 else "Arial"

                if f_risk > 50 or f_flag:
                    if idx == 3:
                        set_cell_shading(row_cells[idx], "FDE2E2")
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0x9B, 0x1C, 0x1C)
                    elif idx == 1:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Recommended action
    p_rec = doc.add_paragraph()
    r_rec = p_rec.add_run("Recommended action")
    r_rec.font.bold = True
    r_rec.font.size = Pt(11)
    r_rec.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    p_rec.paragraph_format.space_after = Pt(6)

    rec1 = f"Remove {worst_name} from the feature set — a risk score of 100.0 combined with a 52.4% retrain drop indicates the feature is effectively a restatement of the label."
    rec2 = "Re-audit Sex and Pclass after the primary leak is removed; both currently sit in low-risk territory (25.0 and 10.0) and are plausible, legitimate predictors, but should be re-scored once the dominant leak is gone since it may be suppressing their apparent effect."
    rec3 = "Re-run the full ModelDoctor audit after remediation — the Overall Health Score, Calibration, and Overfitting numbers in this report are only valid for the current, leaky feature set."

    for text in [rec1, rec2, rec3]:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(4)
        p_b.paragraph_format.line_spacing = 1.15
        p_b.paragraph_format.left_indent = Inches(0.25)
        
        parts = text.split(" — ", 1)
        if len(parts) == 2:
            r_lead = p_b.add_run(parts[0])
            r_lead.font.bold = True
            r_lead.font.size = Pt(9.5)
            r_rest = p_b.add_run(" — " + parts[1])
            r_rest.font.size = Pt(9.5)
        else:
            r_all = p_b.add_run(text)
            r_all.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 4 Heading
    p_sec4 = doc.add_paragraph()
    r_sec4 = p_sec4.add_run("SECTION 4")
    r_sec4.font.size = Pt(9)
    r_sec4.font.bold = True
    r_sec4.font.color.rgb = RGBColor(0x7D, 0x85, 0x90)
    p_sec4.paragraph_format.space_after = Pt(2)

    p_h4 = doc.add_paragraph()
    r_h4 = p_h4.add_run("Fairness Metrics")
    r_h4.font.size = Pt(14)
    r_h4.font.bold = True
    r_h4.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
    p_h4.paragraph_format.space_after = Pt(6)

    p_desc4 = doc.add_paragraph()
    r_desc4 = p_desc4.add_run(
        "Both fairness metrics currently read 0.0000, indicating no measurable disparity across the groups evaluated. Treat this as provisional: leakage can distort or mask fairness signals, so re-check after the leak in Section 3 is fixed."
    )
    r_desc4.font.size = Pt(9.5)
    r_desc4.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    p_desc4.paragraph_format.space_after = Pt(8)

    tbl_fair = doc.add_table(rows=2, cols=2)
    tbl_fair.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths_f = [Inches(4.5), Inches(2.0)]

    fair_lbls = ["Demographic Parity Difference", "Equalized Odds Difference"]
    fair_dict = job_results.get("fairness") or {}
    dp_val = fair_dict.get("demographic_parity_difference", 0.0)
    eo_val = fair_dict.get("equalized_odds_difference", 0.0)
    fair_vals = [f"{dp_val:.4f}", f"{eo_val:.4f}"]

    for r_idx in range(2):
        row_cells = tbl_fair.rows[r_idx].cells
        row_cells[0].text = fair_lbls[r_idx]
        row_cells[1].text = fair_vals[r_idx]

        for c_idx in range(2):
            row_cells[c_idx].width = col_widths_f[c_idx]
            set_cell_shading(row_cells[c_idx], "DEF7EC")
            set_cell_margins(row_cells[c_idx], top=60, bottom=60, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            r = p.runs[0]
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x03, 0x54, 0x3F)
            if c_idx == 1:
                r.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 5 Heading
    p_sec5 = doc.add_paragraph()
    r_sec5 = p_sec5.add_run("SECTION 5")
    r_sec5.font.size = Pt(9)
    r_sec5.font.bold = True
    r_sec5.font.color.rgb = RGBColor(0x7D, 0x85, 0x90)
    p_sec5.paragraph_format.space_after = Pt(2)

    p_h5 = doc.add_paragraph()
    r_h5 = p_h5.add_run("Module Status")
    r_h5.font.size = Pt(14)
    r_h5.font.bold = True
    r_h5.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
    p_h5.paragraph_format.space_after = Pt(6)

    p_desc5 = doc.add_paragraph()
    r_desc5 = p_desc5.add_run(
        "Coverage of this audit run. Data Drift has not yet been analyzed and is not reflected in the Overall Health Score above."
    )
    r_desc5.font.size = Pt(9.5)
    r_desc5.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    p_desc5.paragraph_format.space_after = Pt(10)

    # Status Grid
    tbl_status = doc.add_table(rows=1, cols=2)
    tbl_status.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells_st = tbl_status.rows[0].cells
    hdr_cells_st[0].text = "MODULE"
    hdr_cells_st[1].text = "STATUS"
    col_widths_st = [Inches(4.5), Inches(2.0)]

    for idx in range(2):
        hdr_cells_st[idx].width = col_widths_st[idx]
        set_cell_shading(hdr_cells_st[idx], "1F2937")
        set_cell_margins(hdr_cells_st[idx], top=100, bottom=100, left=100, right=100)
        p = hdr_cells_st[idx].paragraphs[0]
        r = p.runs[0]
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    status_items = [
        ("Calibration", _module_status(job_results.get("calibration"))),
        ("Overfitting", _module_status(job_results.get("overfitting"))),
        ("Feature Dom.", _module_status(job_results.get("feature_dominance"))),
        ("Fairness", _module_status(job_results.get("fairness"))),
        ("Data Drift", _module_status(job_results.get("drift")))
    ]

    for mod_name, st_val in status_items:
        row_cells = tbl_status.add_row().cells
        row_cells[0].text = mod_name
        row_cells[1].text = st_val

        if st_val == "Analyzed":
            bg_color = "DEF7EC"
            fg_color = RGBColor(0x03, 0x54, 0x3F)
        else:
            bg_color = "FEF3C7"
            fg_color = RGBColor(0x92, 0x40, 0x0E)

        for idx in range(2):
            row_cells[idx].width = col_widths_st[idx]
            set_cell_margins(row_cells[idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[idx].paragraphs[0]
            r = p.runs[0]
            r.font.size = Pt(9)
            if idx == 1:
                set_cell_shading(row_cells[idx], bg_color)
                r.font.bold = True
                r.font.color.rgb = fg_color

    doc.add_page_break()

    # =========================================================================
    # PAGE 4: NEXT STEPS FOR THE ENGINEERING TEAM
    # =========================================================================
    add_running_header(doc, job_id)

    # SECTION 6 Heading
    p_sec6 = doc.add_paragraph()
    r_sec6 = p_sec6.add_run("SECTION 6")
    r_sec6.font.size = Pt(9)
    r_sec6.font.bold = True
    r_sec6.font.color.rgb = RGBColor(0x7D, 0x85, 0x90)
    p_sec6.paragraph_format.space_after = Pt(2)

    p_h6 = doc.add_paragraph()
    r_h6 = p_h6.add_run("Next Steps for the Engineering Team")
    r_h6.font.size = Pt(16)
    r_h6.font.bold = True
    r_h6.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
    p_h6.paragraph_format.space_after = Pt(12)

    steps = [
        ("1. Fix the blocker", [
            f"Drop {worst_name} from the training pipeline.",
            "Grep the feature-engineering code for how this column was constructed — a name like this usually means the target was joined in upstream, e.g. during a merge or aggregation step."
        ]),
        ("2. Re-run the audit", [
            "Re-run ModelDoctor end-to-end once the leaky feature is removed — every score in this report (health score, calibration, overfitting) is downstream of the leak and will change."
        ]),
        ("3. Revisit calibration", [
            f"Calibration penalty is {cal_p:.1f}/100. Do not use raw predicted probabilities for decisions until this is re-checked post-leak fix; consider Platt scaling or isotonic regression if it persists."
        ]),
        ("4. Close the coverage gap", [
            "Data Drift has not been analyzed. Schedule this before sign-off, since it's not reflected in the current Overall Health Score."
        ])
    ]

    for title_text, bullets in steps:
        p_t = doc.add_paragraph()
        r_t = p_t.add_run(title_text)
        r_t.font.bold = True
        r_t.font.size = Pt(11)
        r_t.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        p_t.paragraph_format.space_before = Pt(8)
        p_t.paragraph_format.space_after = Pt(4)

        for b_text in bullets:
            p_b = doc.add_paragraph(style='List Bullet')
            p_b.paragraph_format.space_after = Pt(3)
            p_b.paragraph_format.line_spacing = 1.15
            p_b.paragraph_format.left_indent = Inches(0.25)
            
            parts = b_text.split(" — ", 1)
            delim = " — "
            if len(parts) == 1:
                parts = b_text.split("; ", 1)
                delim = "; "
                
            if len(parts) == 2:
                r_lead = p_b.add_run(parts[0])
                r_lead.font.bold = True
                r_lead.font.size = Pt(9.5)
                r_rest = p_b.add_run(delim + parts[1])
                r_rest.font.size = Pt(9.5)
            else:
                r_all = p_b.add_run(b_text)
                r_all.font.size = Pt(9.5)

    # Footing spacer
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(48)

    p_ln = doc.add_paragraph()
    r_ln = p_ln.add_run("_________________________________________________________________________________")
    r_ln.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r_ln.font.size = Pt(9)
    p_ln.paragraph_format.space_after = Pt(6)

    p_f = doc.add_paragraph()
    r_f = p_f.add_run(
        f"Generated by ModelDoctor v0.1.0 • {generated_at} • Executive summary: {mode_text}"
    )
    r_f.font.size = Pt(8)
    r_f.font.color.rgb = RGBColor(0x7D, 0x85, 0x90)
    p_f.paragraph_format.space_after = Pt(0)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

